# w3c.py - convert W3C recommendations in HTML format to a Word file

import os
import re
import sys
import lxml.etree as et
from lxml import objectify

from docx_common import black
from docx_common import add_text_run, new_paragraph, add_heading, add_title
from docx_common import add_internal_link, add_hyperlink, add_list_item

# Writing out to Word .docx files
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT

#-------------------------------------------------------------------------------

def coalesce(s):
    """Coalesce multiple whitespace characters into a single space char."""
    return re.sub(r'\s+', ' ', s)

# -----------------------------------------------------------------------------

def get_tr_size(nd):
    nb_cols = 0
    for k in nd:
        if k.tag in ['th', 'td']:
            span = int(k.attrib['colspan']) if 'colspan' in k.attrib else 1
            if span > nb_cols:
                nb_cols = span
        else:
            m = f'Unexpected tag "{k.tag}" inside a <tbody> element'
            raise RuntimeError(m)

    return nb_cols

# -----------------------------------------------------------------------------

def get_tbody_size(nd):
    # nd is a <tbody> element
    nb_cols_max = 0
    rows = 0
    for k in nd:
        if k.tag == 'tr':
            nb_cols = get_tr_size(k)
            if nb_cols > nb_cols_max:
                nb_cols_max = nb_cols
            rows += 1
        else:
            m = f'Unexpected tag "{k.tag}" inside a <tbody> element'
            raise RuntimeError(m)
    return rows, nb_cols_max

# -----------------------------------------------------------------------------

def get_table_sizes(nd):
    tables = nd.xpath('.//*[self::div1 or self::div2 or self::div3]/table/tbody')
    return [get_tbody_size(t) for t in tables]

# -----------------------------------------------------------------------------

def get_refs(nd):
    refs = {}
    divs = nd.xpath('.//*[self::div1 or self::div2 or self::div3]')
    for k in divs:
        if 'id' in k.attrib:
            id_ = k.attrib['id']
            refs[id_] = k.find('.//head').text
    return refs

# -----------------------------------------------------------------------------

class Docx:
    def do_eg(self, nd):
        # One paragraph for everyting
        p = new_paragraph(self.doc, style='Code')

        if nd.text is not None:
            add_text_run(p, nd.text.strip(), font='Consolas', sz=9)

    #---------------------------------------------------------------------------

    def do_emph(self, nd, p, italic=None):
        # Text before any eventual sub-element.
        if nd.text is not None:
            s = nd.text
            s = coalesce(s)
            if len(s) > 0:
                add_text_run(p, s, italic=italic)

        # Elements under <emph>: phrase, 

        # Process any eventual sub-elements.
        for k in nd:
            if k.tag == 'phrase':
                self.do_phrase(k, p, italic=italic)
            else:
                m = f'Unexpected tag "{k.tag}" inside an <emph> element'
                raise RuntimeError(m)

        if nd.tail:
            s = nd.tail
            s = coalesce(s)
            if len(s) > 0:
                add_text_run(p, s, italic=italic)

    #---------------------------------------------------------------------------

    def do_xspecref(self, nd, p):
        url = nd.attrib['href']
        s = nd.text.strip()
        s = coalesce(s.replace('\n', ' '))
        add_hyperlink(p, s, url)

        if nd.tail:
            s = nd.tail
            s = coalesce(s)
            if len(s) > 0:
                add_text_run(p, s)

    #---------------------------------------------------------------------------

    def do_loc(self, nd, p, bold=None, font=None, sz=None):
        url = nd.attrib['href']
        s = nd.text.strip()
        s = coalesce(s.replace('\n', ' '))
        add_hyperlink(p, s, url, font=font, sz=sz)

        if nd.tail:
            s = nd.tail
            s = coalesce(s)
            if len(s) > 0:
                add_text_run(p, s, font=font, sz=sz)

    #---------------------------------------------------------------------------

    def do_specref(self, nd, p):
        ref = nd.attrib['ref']
        text = self.refs[ref]
        add_internal_link(p, text, text)

        if nd.tail:
            s = nd.tail
            s = coalesce(s)
            if len(s) > 0:
                add_text_run(p, s)

    #---------------------------------------------------------------------------

    def do_phrase(self, nd, p, italic=None):
        # Attributes: diff == add/chg: keep it, del: ignore it
        diff = nd.attrib['diff']
        if diff not in ['chg', 'add', 'del']:
            m = f'Unexpected diff attribute value "{diff}" in a <phrase> element'
            raise RuntimeError(m)

        # Text before any eventual sub-element
        if diff != 'del':
            if nd.text is not None:
                s = nd.text
                s = coalesce(s)
                if len(s) > 0:
                    add_text_run(p, s, italic=italic)

            # Elements under <phrase>: loc, xspecref, code, 

            # Process any eventual sub-elements.
            # FIXME should pass in italic=italic as well
            for k in nd:
                if k.tag == 'loc':
                    self.do_loc(k, p)
                elif k.tag == 'xspecref':
                    self.do_xspecref(k, p)
                elif k.tag == 'code':
                    self.do_code(k, p)
                else:
                    m = f'Unexpected tag "{k.tag}" inside a <div> element'
                    raise RuntimeError(m)

        # Always process tail, regardless of 'diff' value
        if nd.tail:
            s = nd.tail
            s = coalesce(s)
            if len(s) > 0:
                add_text_run(p, s, italic=italic)

    #---------------------------------------------------------------------------

    def do_code(self, nd, p):
        # Text before any eventual sub-element
        if nd.text is not None:
            s = nd.text
            s = coalesce(s.replace('\n', ' ')).strip()
            if len(s) > 0:
                add_text_run(p, s, bold=True, font='Consolas', sz=9)

        # Elements under <code>: loc

        # Process any eventual sub-elements.
        for k in nd:
            if k.tag == 'loc':
                self.do_loc(k, p, font='Consolas', sz=9)
            else:
                m = f'Unexpected tag "{k.tag}" inside a <code> element'
                raise RuntimeError(m)

        if nd.tail:
            s = nd.tail
            s = coalesce(s)
            if len(s) > 0:
                # Outside the <code> element, back to normal font
                add_text_run(p, s)
            
    def do_head(self, nd, level):
        head = nd.text if nd.text is not None else '<empty>'
        add_heading(self.doc, level, head)

    #---------------------------------------------------------------------------

    def do_p(self, nd):
        # One paragraph for everyting (passed onto child elements)
        p = new_paragraph(self.doc)

        # Text before any eventual sub-element.
        if nd.text is not None:
            s = nd.text.lstrip()
            s = coalesce(s)
            if len(s) > 0:
                add_text_run(p, s)

        # Elements under <p>: phrase, xspecref, loc, specref, emph, code

        # Process any eventual sub-elements.
        for k in nd:
            if k.tag == 'phrase':
                self.do_phrase(k, p)
            elif k.tag == 'xspecref':
                self.do_xspecref(k, p)
            elif k.tag == 'loc':
                self.do_loc(k, p)
            elif k.tag == 'specref':
                self.do_specref(k, p)
            elif k.tag == 'emph':
                self.do_emph(k, p, italic=True)
            elif k.tag == 'code':
                self.do_code(k, p)
            else:
                m = f'Unexpected tag "{k.tag}" inside a <p> element'
                raise RuntimeError(m)

        if nd.tail:
            s = nd.tail
            s = coalesce(s)
            if len(s) > 0:
                add_text_run(p, s)

    #---------------------------------------------------------------------------

    def do_note(self, nd):
        role = nd.attrib['role'].capitalize()

        p = new_paragraph(self.doc)
        add_text_run(p, role, font='Arial', sz=10, bold=True)

        for k in nd:
            if k.tag == 'p':
                self.do_p(k)
            elif k.tag == 'eg':
                self.do_eg(k)
            else:
                m = f'Unexpected tag "{k.tag}" inside a <note> element'
                raise RuntimeError(m)

    #---------------------------------------------------------------------------

    def do_cell(self, nd):
        colspan = int(nd.attrib['colspan']) if 'colspan' in nd.attrib else 1

        # Elements under <th>/<td>: table, loc, phrase, xspecref, code, specref
        # All these elements take a paragraph as parameter

        return colspan, nd.text

    #---------------------------------------------------------------------------

    def do_tr(self, nd):
        nb_cols = 0
        row = []
        for k in nd:
            if k.tag in ['th', 'td']:
                span, text = self.do_cell(k)
                colspan = 0
                if span > colspan:
                    colspan = span
                row.append((colspan, text))
                if colspan > nb_cols:
                    nb_cols = colspan
            else:
                m = f'Unexpected tag "{k.tag}" inside a <tbody> element'
                raise RuntimeError(m)

        return nb_cols, row

    #---------------------------------------------------------------------------

    def do_tbody(self, nd):
        nb_cols_max = 0
        rows = []
        for k in nd:
            if k.tag == 'tr':
                nb_cols, row = self.do_tr(k)
                if nb_cols > nb_cols_max:
                    nb_cols_max = nb_cols
                rows.append(row)
            else:
                m = f'Unexpected tag "{k.tag}" inside a <tbody> element'
                raise RuntimeError(m)

        # We have all the data, now create the table
        tbl = self.doc.add_table(len(rows), nb_cols_max)
        tbl.style = 'Table Grid'
        for i, row in enumerate(rows):
            r = tbl.rows[i]
            for j, (colspan, text) in enumerate(row):
                r.cells[j].text = text if text is not None else ''
                if colspan > 1:
                    r.cells[j].merge(r.cells[j+colspan-1])

    #---------------------------------------------------------------------------

    def do_table(self, nd):
        for k in nd:
            if k.tag == 'tbody':
                self.do_tbody(k)
            else:
                m = f'Unexpected tag "{k.tag}" inside a <table> element'
                raise RuntimeError(m)

    # --------------------------------------------------------------------------

    def do_orglist(self, nd):
        # Has attribute diff="add"
        for k in nd:
            if k.tag == 'member':
                name = k.find('.//name').text.strip()
                affil = k.find('.//affiliation').text.strip()
                add_list_item(self.doc, f'{name}, {affil}')

    # --------------------------------------------------------------------------

    def do_div(self, nd):
        level = int(nd.tag[3])  # 1, 2, 3 or 4

        for k in nd:
            if k.tag == 'head':
                self.do_head(k, level)
            elif k.tag == 'p':
                self.do_p(k)
            elif k.tag == 'note':
                self.do_note(k)
            elif k.tag == 'table':
                self.do_table(k)
            elif k.tag == 'orglist':
                self.do_orglist(k)
            elif k.tag in ['div1', 'div2', 'div3']:
                self.do_div(k)
            else:
                m = f'Unexpected tag "{k.tag}" inside a <div> element'
                raise RuntimeError(m)

    #---------------------------------------------------------------------------

    def do_status(self, nd):
        add_title(self.doc, 'Status of this Document')

        # Elements under <status>: p, 

        # Process any eventual sub-elements.
        for k in nd:
            if k.tag == 'p':
                self.do_p(k)
            else:
                m = f'Unexpected tag "{k.tag}" inside an <status> element'
                raise RuntimeError(m)

    #---------------------------------------------------------------------------

    def do_abstract(self, nd):
        # FIXME this is identical to status
        add_title(self.doc, 'Abstract')

        # Elements under <abstract>: p, 

        # Process any eventual sub-elements.
        for k in nd:
            if k.tag == 'p':
                self.do_p(k)
            else:
                m = f'Unexpected tag "{k.tag}" inside an <abstract> element'
                raise RuntimeError(m)

    #---------------------------------------------------------------------------

    def do_header(self, nd):
        """<header> includes <status> and <abstract>."""
        for k in nd:
            if k.tag == 'status':
                self.do_status(k)
            elif k.tag == 'abstract':
                self.do_abstract(k)
            elif k.tag in ['title', 'version', 'w3c-designation',
                            'w3c-doctype', 'pubdate', 'publoc', 'altlocs',
                           'latestloc', 'prevlocs', 'authlist', 'errataloc',
                           'preverrataloc', 'translationloc', 'sourcedesc',
                           'langusage', 'revisiondesc']:
                m = f'Support for tag "{k.tag}" not implemented'
                print(m, file=sys.stderr)
            else:
                m = f'Unexpected tag "{k.tag}" inside a <header> element'
                raise RuntimeError(m)

    #---------------------------------------------------------------------------

    def do_body(self, nd):
        """<body> holds the document content."""
        for k in nd:
            if k.tag in ['div1', 'div2', 'div3']:
                self.do_div(k)
            else:
                m = f'Unexpected tag "{k.tag}" inside a <spec> element'
                raise RuntimeError(m)

    #---------------------------------------------------------------------------

    def do_back(self, nd):
        """<back> holds acknowledgements and appendixes."""
        for k in nd:
            if k.tag in ['div1', 'div2', 'div3']:
                self.do_div(k)
            else:
                m = f'Unexpected tag "{k.tag}" inside a <back> element'
                raise RuntimeError(m)

    #---------------------------------------------------------------------------

    def do_spec(self, nd):
        """<spec> is the toplevel element in the XML tree."""
        for k in nd:
            if k.tag == 'header':
                self.do_header(k)
            elif k.tag == 'body':
                self.do_body(k)
            elif k.tag == 'back':
                self.do_back(k)
            else:
                m = f'Unexpected tag "{k.tag}" inside a <spec> element'
                raise RuntimeError(m)

    #---------------------------------------------------------------------------
    
    def __init__(self, filepath):
        self.filepath = filepath
        
        # Parse XML file, ignoring comments
        p = et.XMLParser(remove_comments=True)
        self.root = objectify.parse(filepath, parser=p).getroot()

        self.refs = get_refs(self.root)
        # for k, v in refs.items():
        #     print(f'{k}: {v}')

        self.tbl_sz = get_table_sizes(self.root)
        # for r, c in tbl_sz:
        #     print(f'({r}, {c})')
        self.curr_table = 0
        
        # Write out the contents as Word
        self.doc = Document('empty.docx')

        # Set normal margins
        s = self.doc.sections[0]
        s.left_margin = Inches(0.59)
        s.right_margin = Inches(0.59)
        s.top_margin = Inches(0.59)
        s.bottom_margin = Inches(0.59)

        # Get the document contents
        self.do_spec(self.root)

    def write(self):
        # Output a .docx file in the current directory
        base, _ = os.path.splitext(self.filepath)
        outpath = f'{base}.docx'
        self.doc.save(outpath)
    
# -----------------------------------------------------------------------------
# main
# -----------------------------------------------------------------------------
    
if __name__ == '__main__':
    # Command line argument
    if len(sys.argv) != 2:
        print(f'Usage: {sys.argv[0]} <xml filepath>')
        exit(-1)
    filepath = sys.argv[1]

    d = Docx(filepath)
    d.write()
