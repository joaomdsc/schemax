# docx_common.py - common functions for the creation of  Word .docx files

import re

# Writing out to Word .docx files
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_BREAK
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE

# Colors
black = RGBColor(0x0, 0x0, 0x0)

# Inside a paragraph, true if the previous output was terminated by a
# non-whitespace character (the paragraph itself is terminated by a newline).
prev_no_space = False
    
#--------------------------------------------------------------------------
# Create a new paragraph object
#--------------------------------------------------------------------------

def new_paragraph(doc, style=None):
    global prev_no_space
    
    p = doc.add_paragraph(style=style)

    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    prev_no_space = False
    
    return p
  
#--------------------------------------------------------------------------
# Add a text run
#--------------------------------------------------------------------------

def add_text_run(p, text, bold=None, italic=None, strike=None, bg=None,
                 font=None, sz=None):
    global prev_no_space
    if text == '':
        return

    # One run for plain text
    r = p.add_run()
    
    f = r.font
    # f.name = 'Times New Roman'
    f.name = 'Calibri' if font is None else font
    f.size = Pt(10 if sz is None else sz)

    r.text = text
    # r.text = ' ' + text if prev_no_space else text
    # prev_no_space = text[-1] not in [' \t\r\n']
    
    if bold:
        r.bold = bold
    if italic:
        r.italic = italic

#--------------------------------------------------------------------------
# Write one heading with style Heading 1, 2, or 3
#--------------------------------------------------------------------------

def add_heading(doc, level, text0):
    # One paragraph for the heading line
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    r = p.add_run()

    r.text = text0

    # Bookmark from https://stackoverflow.com/questions/57586400
    tag = r._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), '0')
    start.set(qn('w:name'), text0)
    tag.append(start)

    text = OxmlElement('w:r')
    text.text = text
    tag.append(text)

    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), '0')
    end.set(qn('w:name'), text0)
    tag.append(end)

#-------------------------------------------------------------------------------
# Add a link to a bookmark
#-------------------------------------------------------------------------------

# Limitations: bookmark names must be <= 40 characters long, with no whitespace
# (replace with undersocres). Punctuation is not allowed, but in practice the
# colon ':' is accepted.

def add_internal_link(p, link_to, text):
    link_to = link_to[:40].replace(' ', '_')
    
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('w:anchor'), link_to)

    r = OxmlElement('w:r')
    r.append(OxmlElement('w:rPr'))
    r.text = text
    hl.append(r)

    r = p.add_run()
    r._r.append(hl)
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

#-------------------------------------------------------------------------------
# Add an external hyperlink
#-------------------------------------------------------------------------------

def add_hyperlink(p, text, url, font='Calibri', sz=10):
    part = p.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    # Create the w:hyperlink tag and add needed values
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)

    # Create a w:r element and a new w:rPr element. Join all the xml elements
    # together and add the required text to the w:r element
    r = OxmlElement('w:r')
    r.append(OxmlElement('w:rPr'))
    r.text = text
    hl.append(r)

    r = p.add_run()
    r._r.append(hl)
    r.font.name = 'Calibri' if font is None else font
    r.font.size = Pt(10 if sz is None else sz)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
   
# -----------------------------------------------------------------------------
# main
# -----------------------------------------------------------------------------
    
if __name__ == '__main__':
    print('This module is not meant to be executed directly')