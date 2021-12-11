# msword.py - MS Word .docx file

# Writing out to Word .docx files
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_LINE_SPACING
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE

#--------------------------------------------------------------------------

class MsPara:
    def __init__(self, doc, style=None):
        self.p = doc.add_paragraph(style=style)
        self.pending_bookmark = None

        pf = self.p.paragraph_format
        pf.space_before = Pt(4)
        pf.space_after = Pt(4)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    #----------------------------------------------------------------------

    def add_text_run(self, text, bold=None, italic=None, strike=None,
                     font=None, sz=None):
        if text == '':
            return

        # One run for plain text
        r = self.p.add_run()
        f = r.font
        f.name = 'Calibri' if font is None else font
        f.size = Pt(10 if sz is None else sz)
        r.text = text
        if bold:
            r.bold = bold
        if italic:
            r.italic = italic

        # Bookmarking a paragraph requires a text run
        if self.pending_bookmark is not None:
            Docx.add_bookmark(r, self.pending_bookmark)
            self.pending_bookmark = None

    #---------------------------------------------------------------------------

    def add_internal_link(self, link_to, text):
        """Add a link to a bookmark.

        Limitations: bookmark names must be <= 40 characters long, with no
        whitespace (replace with underscores). Punctuation is not allowed, but
        in practice the colon ':' is accepted.

        """
        link_to = link_to[:40].replace(' ', '_')

        hl = OxmlElement('w:hyperlink')
        hl.set(qn('w:anchor'), link_to)

        wr = OxmlElement('w:r')
        wr.append(OxmlElement('w:rPr'))
        wr.text = text
        hl.append(wr)

        r = self.p.add_run()
        r._r.append(hl)
        r.font.name = "Calibri"
        r.font.size = Pt(10)
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True

        # Bookmarking a paragraph requires a text run
        if self.pending_bookmark is not None:
            Docx.add_bookmark(r, self.pending_bookmark)
            self.pending_bookmark = None

    #---------------------------------------------------------------------------

    def add_hyperlink(self, text, url, font='Calibri', sz=10):
        """Add an external hyperlink."""
        part = self.p.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK,
                              is_external=True)

        # Create the w:hyperlink tag and add needed values
        hl = OxmlElement('w:hyperlink')
        hl.set(qn('r:id'), r_id)

        # Create a w:r element and a new w:rPr element. Join all the xml
        # elements together and add the required text to the w:r element
        wr = OxmlElement('w:r')
        wr.append(OxmlElement('w:rPr'))
        wr.text = text
        hl.append(wr)

        r = self.p.add_run()
        r._r.append(hl)
        r.font.name = 'Calibri' if font is None else font
        r.font.size = Pt(10 if sz is None else sz)
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True

        # Bookmarking a paragraph requires a text run
        if self.pending_bookmark is not None:
            Docx.add_bookmark(r, self.pending_bookmark)
            self.pending_bookmark = None

class MsCellPara(MsPara):
    def __init__(self, cell):
        self.p = cell.add_paragraph()
        self.pending_bookmark = None

        pf = self.p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

#-------------------------------------------------------------------------------

class Table:
    def __init__(self, doc, nrows, ncols):
        """Add a table with 'nrows' rows and 'ncols' columns."""
        self.tbl = doc.add_table(nrows, ncols)
        self.tbl.style = 'Table Grid'

    def get_row(self, idx):
        return self.tbl.rows[idx]

#-------------------------------------------------------------------------------

class Docx:
    def __init__(self):
        self.doc = Document('empty.docx')

        # Set normal margins
        s = self.doc.sections[0]
        s.left_margin = Inches(0.59)
        s.right_margin = Inches(0.59)
        s.top_margin = Inches(0.59)
        s.bottom_margin = Inches(0.59)
        
    def write(self, filepath):
        self.doc.save(filepath)

    def new_paragraph(self, style=None):
        return MsPara(self.doc, style=style)

    #---------------------------------------------------------------------------

    def add_heading(self, level, text):
        """Write one heading with style Heading 1, 2, or 3."""
        # One paragraph for the heading line
        p = self.doc.add_paragraph()
        p.style = self.doc.styles[f'Heading {level}']

        r = p.add_run()
        r.text = text
        Docx.add_bookmark(r, text)

    #---------------------------------------------------------------------------

    def add_title(self, text):
        """Write a title without making it a header (without numbering)."""
        # One paragraph for the heading line
        p = self.doc.add_paragraph()
        r = p.add_run()

        f = r.font
        f.name = 'Arial'
        f.size = Pt(12)
        r.bold = True

        r.text = text

    #---------------------------------------------------------------------------

    def add_list_item(self, text):
        """Add an item in an unordered list."""
        self.doc.add_paragraph(text, style='List Paragraph')

    #---------------------------------------------------------------------------

    def new_table(self, nrows, ncols):
        """Add a table with 'nrows' rows and 'ncols' columns."""
        return Table(self.doc, nrows, ncols)

    #---------------------------------------------------------------------------

    @staticmethod
    def add_bookmark(r, text):
        """Add bookmark to a text run"""
        # Bookmark from https://stackoverflow.com/questions/57586400

        tag = r._r
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), '0')
        start.set(qn('w:name'), text)
        tag.append(start)

        wr = OxmlElement('w:r')
        wr.text = ''
        tag.append(wr)

        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), '0')
        end.set(qn('w:name'), text)
        tag.append(end)

# -----------------------------------------------------------------------------
# main
# -----------------------------------------------------------------------------
    
if __name__ == '__main__':
    print('This module is not meant to be executed directly')